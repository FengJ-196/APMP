#!/usr/bin/env python3
"""
modify_docx.py - A comprehensive utility script to inspect and modify Microsoft Word (.docx) documents.
Uses python-docx to perform search-and-replace (preserving formatting), list headings,
inspect tables, and insert new elements.

Prerequisites:
  pip install python-docx

Usage:
  python scripts/modify_docx.py --help
"""

import os
import sys
import argparse
from docx import Document

# Reconfigure stdout and stderr to use UTF-8 to support Vietnamese characters in Windows console
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def get_doc_path(path_arg):
    """Resolve docx path, default to the SDD file if not specified or not found."""
    default_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "docs",
        "Tai lieu thiet ke phan mem SDD (1).docx"
    )
    
    if path_arg:
        if os.path.exists(path_arg):
            return path_arg
        # Try relative to current script
        alt_path = os.path.abspath(path_arg)
        if os.path.exists(alt_path):
            return alt_path
        print(f"Warning: Specified file not found at '{path_arg}'. Falling back to default.")
    
    if os.path.exists(default_path):
        return default_path
    
    # Fallback to similar named file in docs/
    docs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
    if os.path.exists(docs_dir):
        files = os.listdir(docs_dir)
        docx_files = [f for f in files if f.endswith('.docx')]
        if docx_files:
            # Prefer SDD file
            sdd_files = [f for f in docx_files if "sdd" in f.lower() or "thiet ke" in f.lower()]
            if sdd_files:
                return os.path.join(docs_dir, sdd_files[0])
            return os.path.join(docs_dir, docx_files[0])
            
    print(f"Error: Could not locate default Word document. Please provide a path using --file.")
    sys.exit(1)

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Replaces old_text with new_text in a paragraph while preserving formatting/styles.
    Maps characters of paragraph.text to individual runs and updates only the necessary runs.
    """
    if old_text not in paragraph.text:
        return False
    
    # Simple case: check if old_text is contained fully inside a single run
    replaced = False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True
            
    # If we successfully replaced it and the term is no longer in the paragraph text, we are done
    if replaced and old_text not in paragraph.text:
        return True
            
    # Complex case: term is split across multiple runs
    # Map text indexes to run objects
    text = ""
    run_ranges = []
    for run in paragraph.runs:
        start = len(text)
        text += run.text
        end = len(text)
        run_ranges.append((start, end, run))
        
    start_idx = 0
    paragraph_modified = False
    while True:
        idx = text.find(old_text, start_idx)
        if idx == -1:
            break
            
        end_idx = idx + len(old_text)
        overlapping_runs = []
        for start, end, run in run_ranges:
            if start < end_idx and end > idx:
                overlapping_runs.append((start, end, run))
                
        if overlapping_runs:
            first_start, first_end, first_run = overlapping_runs[0]
            prefix = first_run.text[:idx - first_start]
            
            last_start, last_end, last_run = overlapping_runs[-1]
            suffix = last_run.text[end_idx - last_start:]
            
            # Put replacement in the first run
            first_run.text = prefix + new_text
            
            # Clear intermediate runs
            for _, _, run in overlapping_runs[1:-1]:
                run.text = ""
                
            # Put suffix in the last run (if it's different from the first run)
            if len(overlapping_runs) > 1:
                overlapping_runs[-1][2].text = suffix
            else:
                first_run.text += suffix
                
            paragraph_modified = True
            
            # Reconstruct ranges since text structure changed
            text = ""
            run_ranges = []
            for run in paragraph.runs:
                start = len(text)
                text += run.text
                end = len(text)
                run_ranges.append((start, end, run))
                
            start_idx = idx + len(new_text)
        else:
            start_idx = idx + 1
            
    return paragraph_modified or replaced

def replace_text_in_table(table, old_text, new_text):
    """Search and replace text in all cells of a table."""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            # Replace in paragraphs inside cell
            for para in cell.paragraphs:
                if replace_text_in_paragraph(para, old_text, new_text):
                    replaced = True
            # Replace in nested tables inside cell
            for nested_table in cell.tables:
                if replace_text_in_table(nested_table, old_text, new_text):
                    replaced = True
    return replaced

def replace_text_in_document(doc, old_text, new_text):
    """Replaces old_text with new_text in all paragraphs, tables, and headers/footers."""
    para_count = 0
    table_count = 0
    header_footer_count = 0
    
    # 1. Process standard paragraphs
    for para in doc.paragraphs:
        if replace_text_in_paragraph(para, old_text, new_text):
            para_count += 1
            
    # 2. Process tables
    for table in doc.tables:
        if replace_text_in_table(table, old_text, new_text):
            table_count += 1
            
    # 3. Process headers and footers
    for section in doc.sections:
        # Safely collect headers/footers to support different library versions
        header_footers = []
        for attr in ['header', 'first_page_header', 'even_page_header',
                     'footer', 'first_page_footer', 'even_page_footer']:
            if hasattr(section, attr):
                hf = getattr(section, attr)
                if hf:
                    header_footers.append(hf)
                    
        for hf in header_footers:
            for para in hf.paragraphs:
                if replace_text_in_paragraph(para, old_text, new_text):
                    header_footer_count += 1
            for table in hf.tables:
                if replace_text_in_table(table, old_text, new_text):
                    header_footer_count += 1
                    
    return para_count, table_count, header_footer_count

def list_headings(doc):
    """Extract and print the outline of the document based on paragraph styles."""
    print("\n--- Document Outline ---")
    heading_count = 0
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name
        if style_name.startswith('Heading') or style_name == 'Title' or style_name == 'Subtitle':
            indent = "  " * int(style_name.replace('Heading ', '')) if 'Heading ' in style_name else ""
            print(f"[{idx:4d}] {indent}* {para.text} ({style_name})")
            heading_count += 1
    if heading_count == 0:
        print("No standard headings found in the document paragraphs. Printing first 15 paragraphs instead:")
        for idx, para in enumerate(doc.paragraphs[:15]):
            if para.text.strip():
                print(f"[{idx:4d}] ({para.style.name}): {para.text[:80]}...")

def search_text(doc, query):
    """Search the document for a query and print matching paragraphs/cells with context."""
    print(f"\n--- Searching for '{query}' ---")
    matches = 0
    
    # Search in main paragraphs
    for idx, para in enumerate(doc.paragraphs):
        if query.lower() in para.text.lower():
            print(f"[Paragraph {idx}] ({para.style.name}):")
            print(f"  {para.text.strip()}")
            print("-" * 40)
            matches += 1
            
    # Search in tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if query.lower() in cell.text.lower():
                    print(f"[Table {t_idx}, Row {r_idx}, Col {c_idx}]:")
                    print(f"  {cell.text.strip()}")
                    print("-" * 40)
                    matches += 1
                    
    print(f"Found {matches} matches.")

def print_tables_summary(doc):
    """List all tables with their size and partial contents."""
    print("\n--- Tables in Document ---")
    if not doc.tables:
        print("No tables found in this document.")
        return
        
    for idx, table in enumerate(doc.tables):
        num_rows = len(table.rows)
        num_cols = len(table.columns) if num_rows > 0 else 0
        
        # Get preview of headers / first row
        headers = []
        if num_rows > 0:
            headers = [cell.text.strip().replace('\n', ' ')[:20] for cell in table.rows[0].cells[:4]]
            if len(table.rows[0].cells) > 4:
                headers.append("...")
        
        print(f"Table #{idx}: {num_rows} rows x {num_cols} columns")
        print(f"  First Row Preview: {', '.join(headers)}")
        print("-" * 40)

def show_table_detail(doc, table_idx):
    """Print the entire contents of a specific table in a formatted way."""
    if table_idx >= len(doc.tables) or table_idx < 0:
        print(f"Error: Table index {table_idx} is out of range. Total tables: {len(doc.tables)}")
        return
        
    table = doc.tables[table_idx]
    print(f"\n--- Table #{table_idx} Detail ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        row_content = []
        for cell in row.cells:
            # Clean text for single-line display in CLI
            clean_cell = cell.text.strip().replace('\n', ' | ')
            row_content.append(clean_cell)
        print(f"Row {r_idx:2d}: " + "  [AND]  ".join(f"'{col}'" for col in row_content))

def insert_element_after_paragraph(doc, target_para_idx, text, element_type="paragraph", style=None):
    """
    Inserts a new paragraph or heading after a specific paragraph index.
    Using XML manipulation to insert it at the correct position.
    """
    if target_para_idx < 0 or target_para_idx >= len(doc.paragraphs):
        print(f"Error: Paragraph index {target_para_idx} is out of range. Total paragraphs: {len(doc.paragraphs)}")
        return False
        
    target_para = doc.paragraphs[target_para_idx]
    
    if element_type == "heading":
        level = 1
        if style and style.startswith("Heading "):
            try:
                level = int(style.split(" ")[1])
            except ValueError:
                pass
        new_para = doc.add_heading(text, level=level)
    else:
        new_para = doc.add_paragraph(text, style=style)
        
    # Move the new paragraph immediately after the target paragraph in XML structure
    target_para._p.addnext(new_para._p)
    return True

def main():
    parser = argparse.ArgumentParser(description="Modify and inspect Word (.docx) documents.")
    parser.add_argument("-f", "--file", help="Path to the docx file (defaults to Docs/Tai lieu thiet ke phan mem SDD (1).docx)")
    parser.add_argument("-o", "--output", help="Path to save the modified document. If not specified, overrides the input file.")
    
    subparsers = parser.add_subparsers(dest="command", help="Command to run")
    
    # 1. Outline Command
    subparsers.add_parser("outline", help="List all headings in the document (the outline).")
    
    # 2. Search Command
    search_parser = subparsers.add_parser("search", help="Search for a text query in the document.")
    search_parser.add_argument("query", help="Text to search for")
    
    # 3. Replace Command
    replace_parser = subparsers.add_parser("replace", help="Replace text throughout the document.")
    replace_parser.add_argument("old_text", help="Text to search for")
    replace_parser.add_argument("new_text", help="Replacement text")
    replace_parser.add_argument("--dry-run", action="store_true", help="Print matches without saving changes")
    
    # 4. Tables Command
    tables_parser = subparsers.add_parser("tables", help="Inspect tables in the document.")
    tables_parser.add_argument("--detail", type=int, help="Specify table index to show full content")
    
    # 5. Insert Command
    insert_parser = subparsers.add_parser("insert", help="Insert a new paragraph/heading after a specific paragraph index.")
    insert_parser.add_argument("para_idx", type=int, help="Paragraph index to insert after (get this from 'outline' or 'search')")
    insert_parser.add_argument("text", help="Text content of the new paragraph")
    insert_parser.add_argument("--type", choices=["paragraph", "heading"], default="paragraph", help="Type of element (paragraph or heading)")
    insert_parser.add_argument("--style", help="Name of style to apply (e.g. 'Heading 1', 'List Bullet', etc.)")

    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        sys.exit(0)
        
    doc_path = get_doc_path(args.file)
    print(f"Loading document: {doc_path} ...")
    
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        print("Make sure python-docx is installed and the file is not open in Microsoft Word.")
        sys.exit(1)
        
    save_required = False
    
    if args.command == "outline":
        list_headings(doc)
        
    elif args.command == "search":
        search_text(doc, args.query)
        
    elif args.command == "replace":
        print(f"Replacing '{args.old_text}' with '{args.new_text}'...")
        if args.dry_run:
            print("[Dry Run] Finding matches:")
            search_text(doc, args.old_text)
        else:
            p_cnt, t_cnt, hf_cnt = replace_text_in_document(doc, args.old_text, args.new_text)
            print(f"Replacement complete:")
            print(f"  - Modified {p_cnt} paragraphs")
            print(f"  - Modified {t_cnt} tables")
            print(f"  - Modified {hf_cnt} headers/footers")
            if p_cnt > 0 or t_cnt > 0 or hf_cnt > 0:
                save_required = True
                
    elif args.command == "tables":
        if args.detail is not None:
            show_table_detail(doc, args.detail)
        else:
            print_tables_summary(doc)
            
    elif args.command == "insert":
        print(f"Inserting {args.type} after paragraph index {args.para_idx}...")
        success = insert_element_after_paragraph(
            doc, 
            args.para_idx, 
            args.text, 
            element_type=args.type, 
            style=args.style
        )
        if success:
            print("Element successfully inserted.")
            save_required = True
            
    if save_required:
        out_path = args.output if args.output else doc_path
        print(f"Saving changes to: {out_path} ...")
        try:
            doc.save(out_path)
            print("Document saved successfully!")
        except Exception as e:
            print(f"Error saving document: {e}")
            print("Please check if the file is open in another program (like Word) and close it.")

if __name__ == "__main__":
    main()
