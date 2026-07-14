import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys
import re

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def insert_paragraph_before_table(doc, table, text, bold=True, font_size_pt=12):
    """Inserts a paragraph immediately before the given table in the document XML."""
    tbl_element = table._tbl
    parent = tbl_element.getparent()
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = docx.shared.Pt(12)
    p.paragraph_format.space_after = docx.shared.Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = 0 # Left align for table captions
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(font_size_pt)
    run.bold = bold
    
    parent.insert(parent.index(tbl_element), p._p)
    return p

def insert_para_after(doc, target_para, text, bold=False, italic=False, font_size_pt=12, alignment=0, space_before=6, space_after=6):
    """Inserts a paragraph after the target paragraph in the document XML."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = docx.shared.Pt(space_before)
    p.paragraph_format.space_after = docx.shared.Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = alignment
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    
    target_para._p.addnext(p._p)
    return p

def main():
    doc_path = os.path.join("docs", "Mẫu báo cáo cuối kỳ Project, GR.docx")
    if not os.path.exists(doc_path):
        print(f"Error: Could not find document at {doc_path}")
        sys.exit(1)
        
    print(f"Loading document: {doc_path} ...")
    doc = docx.Document(doc_path)
    
    # 1. Clean up any existing duplicate captions in the document first
    print("\n--- Self-Healing: Cleaning up duplicate captions ---")
    cap1_pattern = "so sánh tính năng apmp với các công cụ trợ lý ai hiện nay"
    cap2_pattern = "so sánh khả năng hỗ trợ khởi tạo dự án giữa các giải pháp"
    
    paras_to_remove = []
    found_cap1 = False
    found_cap2 = False
    
    for idx, para in enumerate(doc.paragraphs):
        txt_lower = para.text.strip().lower()
        if cap1_pattern in txt_lower:
            if found_cap1:
                paras_to_remove.append(para)
            else:
                found_cap1 = True
        elif cap2_pattern in txt_lower:
            if found_cap2:
                paras_to_remove.append(para)
            else:
                found_cap2 = True
                
    if paras_to_remove:
        print(f"Removing {len(paras_to_remove)} duplicate caption paragraphs...")
        for p in paras_to_remove:
            p_element = p._p
            p_element.getparent().remove(p_element)
        doc.save(doc_path)
        doc = docx.Document(doc_path)
        
    # 2. Insert captions if not present
    table3 = None
    table4 = None
    for idx, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.columns) == 4:
            c0 = table.rows[0].cells[0].text.strip().lower()
            if "tiêu chí" == c0 and table3 is None:
                table3 = table
            elif "tiêu chí so sánh" == c0 and table4 is None:
                table4 = table
                
    if table3 is None or table4 is None:
        print("Error: Could not locate competitor comparison tables in Section 2.1.")
        sys.exit(1)
        
    # Check if they already have captions
    has_cap1 = False
    has_cap2 = False
    for para in doc.paragraphs:
        txt_lower = para.text.strip().lower()
        if cap1_pattern in txt_lower:
            has_cap1 = True
        if cap2_pattern in txt_lower:
            has_cap2 = True
            
    if not has_cap1:
        print("Inserting caption for Table 3...")
        insert_paragraph_before_table(doc, table3, "Bảng 1: So sánh tính năng APMP với các công cụ trợ lý AI hiện nay", bold=True, font_size_pt=12)
    if not has_cap2:
        print("Inserting caption for Table 4...")
        insert_paragraph_before_table(doc, table4, "Bảng 2: So sánh khả năng hỗ trợ khởi tạo dự án giữa các giải pháp", bold=True, font_size_pt=12)
        
    if not has_cap1 or not has_cap2:
        doc.save(doc_path)
        doc = docx.Document(doc_path)
        
    # 3. Locate 'Giới thiệu đề tài' to define the start of body
    start_body_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if "giới thiệu đề tài" == para.text.strip().lower():
            start_body_idx = idx
            break
            
    if start_body_idx is None:
        print("Error: Could not locate Chapter 1 start heading 'Giới thiệu đề tài'")
        sys.exit(1)
        
    print(f"Document body starts at paragraph index {start_body_idx}.")
    
    # 4. Scan and renumber all figures and tables in the document body
    print("\n--- Scanning and Renumbering Figures and Tables ---")
    fig_pattern = re.compile(r"^Hình\s+(\d+|x|vẽ)[\s\:\-\.]*", re.IGNORECASE)
    tbl_pattern = re.compile(r"^Bảng\s+(\d+|x|vẽ|[\d\.]+)[\s\:\-\.]*", re.IGNORECASE)
    
    fig_list = []
    tbl_list = []
    
    fig_counter = 1
    tbl_counter = 1
    
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip().replace('\xa0', ' ').replace('\u200b', '')
        
        # Check Table Caption
        tbl_match = tbl_pattern.match(txt)
        if tbl_match:
            prefix = tbl_match.group(0)
            clean_text = txt[len(prefix):].strip()
            
            # Update paragraph text in body
            para.text = ""
            para.alignment = 0 # Left align for tables
            para.paragraph_format.space_before = docx.shared.Pt(12)
            para.paragraph_format.space_after = docx.shared.Pt(6)
            para.paragraph_format.line_spacing = 1.15
            
            run_bold = para.add_run(f"Bảng {tbl_counter}: ")
            run_bold.font.name = 'Times New Roman'
            run_bold.font.size = docx.shared.Pt(12)
            run_bold.bold = True
            
            run_normal = para.add_run(clean_text)
            run_normal.font.name = 'Times New Roman'
            run_normal.font.size = docx.shared.Pt(12)
            
            if idx >= start_body_idx:
                tbl_list.append(f"Bảng {tbl_counter}: {clean_text}")
                print(f"  Table {tbl_counter}: {clean_text[:65]}...")
                tbl_counter += 1
                
        # Check Figure Caption
        fig_match = fig_pattern.match(txt)
        if fig_match:
            prefix = fig_match.group(0)
            clean_text = txt[len(prefix):].strip()
            
            # Update paragraph text in body
            para.text = ""
            para.alignment = 1 # Center align for figures
            para.paragraph_format.space_before = docx.shared.Pt(3)
            para.paragraph_format.space_after = docx.shared.Pt(12)
            para.paragraph_format.line_spacing = 1.15
            
            run_bold = para.add_run(f"Hình {fig_counter}: ")
            run_bold.font.name = 'Times New Roman'
            run_bold.font.size = docx.shared.Pt(12)
            run_bold.bold = True
            
            run_normal = para.add_run(clean_text)
            run_normal.font.name = 'Times New Roman'
            run_normal.font.size = docx.shared.Pt(12)
            
            if idx >= start_body_idx:
                fig_list.append(f"Hình {fig_counter}: {clean_text}")
                print(f"  Figure {fig_counter}: {clean_text[:65]}...")
                fig_counter += 1
                
    # Save and reload before writing lists
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    
    # 5. Locate boundaries of the lists in the front matter
    fig_start = None
    fig_end = None
    tbl_start = None
    tbl_end = None
    
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip().lower().replace('\xa0', ' ').replace('\u200b', '')
        if "danh mục hình vẽ" == txt:
            fig_start = idx
        elif "danh mục bảng" == txt:
            fig_end = idx
            tbl_start = idx
        elif "danh mục các từ viết tắt" == txt or "danh mục từ viết tắt" in txt:
            tbl_end = idx
            break
            
    if fig_start is None or fig_end is None or tbl_start is None or tbl_end is None:
        print(f"Error: Could not find list boundaries (fig_start: {fig_start}, fig_end: {fig_end}, tbl_start: {tbl_start}, tbl_end: {tbl_end})")
        sys.exit(1)
        
    print(f"\nWriting List of Figures (between index {fig_start} and {fig_end})...")
    # Delete old List of Figures paragraphs
    paras_to_delete = [doc.paragraphs[k] for k in range(fig_start + 1, fig_end)]
    for p in paras_to_delete:
        p_element = p._p
        p_element.getparent().remove(p_element)
        
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    
    # Re-locate fig_start after deletion
    for idx, para in enumerate(doc.paragraphs):
        if "danh mục hình vẽ" == para.text.strip().lower():
            fig_start = idx
            break
            
    # Insert new List of Figures
    cursor = doc.paragraphs[fig_start]
    for fig_caption in fig_list:
        cursor = insert_para_after(doc, cursor, fig_caption, bold=False, font_size_pt=11, space_before=2, space_after=2)
        
    # Save and reload
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    
    # Re-locate List of Tables boundaries
    tbl_start = None
    tbl_end = None
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip().lower().replace('\xa0', ' ').replace('\u200b', '')
        if "danh mục bảng" == txt:
            tbl_start = idx
        elif "danh mục các từ viết tắt" in txt or "danh mục từ viết tắt" in txt:
            tbl_end = idx
            break
            
    print(f"Writing List of Tables (between index {tbl_start} and {tbl_end})...")
    # Delete old List of Tables paragraphs
    paras_to_delete = [doc.paragraphs[k] for k in range(tbl_start + 1, tbl_end)]
    for p in paras_to_delete:
        p_element = p._p
        p_element.getparent().remove(p_element)
        
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    
    # Re-locate tbl_start after deletion
    for idx, para in enumerate(doc.paragraphs):
        if "danh mục bảng" == para.text.strip().lower():
            tbl_start = idx
            break
            
    # Insert new List of Tables
    cursor = doc.paragraphs[tbl_start]
    for tbl_caption in tbl_list:
        cursor = insert_para_after(doc, cursor, tbl_caption, bold=False, font_size_pt=11, space_before=2, space_after=2)
        
    doc.save(doc_path)
    print("\nList of Figures and List of Tables compiled and updated successfully!")

if __name__ == "__main__":
    main()
