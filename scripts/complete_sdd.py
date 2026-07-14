import docx
import os
import sys

# Reconfigure stdout to use UTF-8 to handle Vietnamese characters properly
try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def add_code_block(doc, target_para, code_lines):
    """Inserts a list of code lines as a Courier New formatted block after target_para."""
    current_p = target_para
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = docx.shared.Inches(0.4)
        p.paragraph_format.space_after = docx.shared.Pt(0)
        p.paragraph_format.space_before = docx.shared.Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = docx.shared.Pt(8.5)
        current_p._p.addnext(p._p)
        current_p = p
    return current_p

def insert_text_paragraphs(doc, target_para, text_list, style='Normal'):
    """Inserts multiple text paragraphs after target_para in sequential order."""
    current_p = target_para
    for text in text_list:
        p = doc.add_paragraph(text, style=style)
        current_p._p.addnext(p._p)
        current_p = p
    return current_p

def main():
    doc_path = os.path.join("docs", "Tai lieu thiet ke phan mem SDD (1).docx")
    if not os.path.exists(doc_path):
        print(f"Error: Could not find document at {doc_path}")
        sys.exit(1)
        
    print(f"Loading document: {doc_path} ...")
    doc = docx.Document(doc_path)
    
    # ----------------------------------------------------
    # TASK 1: Compile Architectural Explanations
    # ----------------------------------------------------
    print("Applying Task 1: Architectural Explanations...")
    
    # Fill in "Lý do lựa chọn:"
    lý_do_text = [
        "1. Kiến trúc phân tầng (Layered Architecture): Tách biệt rõ ràng trách nhiệm giữa giao diện (Presentation), định tuyến bảo mật (API Router), nghiệp vụ cốt lõi (Service), và lưu trữ (Data Layer/Database). Điều này giúp hệ thống dễ dàng bảo trì, phát triển độc lập và viết kiểm thử đơn vị (Unit Test) cho từng service.",
        "2. Tiếp cận Modular Monolith: Đóng gói toàn bộ ứng dụng trong một Next.js codebase giúp tối giản hóa quy trình deploy và quản lý cơ sở hạ tầng trong giai đoạn đầu. Tuy nhiên, các tích hợp bên ngoài (Jira, GitHub, OpenRouter AI) được thiết kế dưới dạng các service module biệt lập, giúp dễ dàng bóc tách ra các microservices khi tải trọng tăng cao.",
        "3. Tối ưu hoá phản hồi Client-side với Zustand Store: Quản lý trạng thái WBS tree phức tạp phía Client một cách reactive mà không bị re-render không cần thiết như các thư viện quản lý state khác, đảm bảo tính mượt mà khi tương tác kéo thả hoặc cập nhật.",
        "4. Tìm kiếm tương tự RAG với Qdrant: Sử dụng vector database chuyên dụng Qdrant chạy độc lập với MongoDB để lưu trữ và tìm kiếm vector (768 chiều sinh bởi Gemini-embedding-2). Việc này giúp giảm tải hoàn toàn cho MongoDB đối với các truy vấn tìm kiếm tương đồng trên hàng chục ngàn issue lịch sử."
    ]
    
    for para in doc.paragraphs:
        if para.text.strip() == "Lý do lựa chọn:":
            # We keep "Lý do lựa chọn:" and insert the items after it
            insert_text_paragraphs(doc, para, lý_do_text)
            break

    # Add explanation for Hình 1: Ảnh kiến trúc phần mềm tổng quan
    hinh1_desc = [
        "Thuyết minh Hình 1: Mô hình kiến trúc thể hiện luồng xử lý từ trình duyệt người dùng đến Next.js App Router Backend qua kết nối HTTPS bảo mật. Khi request đi vào, Auth Middleware sẽ kiểm chuẩn JWT token được lưu ở HTTP-only cookie. Tiếp đến, Route Handler tiếp nhận và chuyển giao tham số tới Service tương ứng (ví dụ: ProjectService, WBSItemService, EstimationService). Để thực hiện RAG, EstimationService gọi QdrantService để sinh vector embedding từ text qua OpenRouter API và tìm kiếm tương đồng trên Qdrant DB. Sau đó, nó đính kèm kết quả lịch sử làm Reference Context và gửi yêu cầu hoàn thành (chat completion) qua AIService đến OpenRouter AI Gateway để đưa ra điểm Story Point tối ưu."
    ]
    for para in doc.paragraphs:
        if para.text.strip() == "Hình 1: Ảnh kiến trúc phần mềm tổng quan":
            insert_text_paragraphs(doc, para, hinh1_desc)
            break

    # Add explanation for Hình 2: Ảnh thiết kế gói tổng quan
    hinh2_desc = [
        "Thuyết minh Hình 2: Sơ đồ gói thể hiện tính cô lập cao giữa các tầng (Layer Boundaries). Tầng Presentation sử dụng API Client để giao tiếp với API Layer. API Layer chuyển hướng yêu cầu nghiệp vụ đến Service Layer. Service Layer phụ thuộc vào Data Layer để CRUD dữ liệu MongoDB (thông qua Mongoose Schema) và Qdrant Vector DB, đồng thời gọi các External API OAuth (Jira REST API, GitHub REST API) để đồng bộ backlog. Nguyên tắc bất di bất dịch là tầng trên chỉ phụ thuộc tầng dưới kế cận, không có dependency ngược hoặc vòng lặp chéo."
    ]
    for para in doc.paragraphs:
        if para.text.strip() == "Hình 2: Ảnh thiết kế gói tổng quan":
            insert_text_paragraphs(doc, para, hinh2_desc)
            break

    # Add explanation for Hình 3: Ảnh gói giao diện và trạng thái
    hinh3_desc = [
        "Thuyết minh Hình 3: Gói giao diện tổ chức xung quanh ProjectDashboard, là container quản lý các tabs như SourceOfTruthPanel, ConflictPanel, WBSPanel, và IntegrationsModal. Trạng thái dự án và dữ liệu WBS được lưu trữ tập trung ở Client-side Zustand Store (useProjectStore). Giao diện UI lắng nghe trạng thái của Store để tự động render. Khi có cập nhật từ người dùng, Store gọi APIClient để đẩy request tới Next.js Backend và cập nhật lại state cục bộ sau khi nhận kết quả thành công."
    ]
    for para in doc.paragraphs:
        if para.text.strip() == "Hình 3: Ảnh gói giao diện và trạng thái":
            insert_text_paragraphs(doc, para, hinh3_desc)
            break

    # Add explanation for Hình 4: Ảnh gói API và bảo mật
    hinh4_desc = [
        "Thuyết minh Hình 4: API Layer chứa các Router và Route Handlers của Next.js tương ứng với các phân hệ nghiệp vụ. Tất cả các router bảo mật (như /projects, /wbs, /estimate) đều có quan hệ phụ thuộc chặt chẽ vào AuthMiddleware. Middleware này thực hiện giải mã token JWT, nếu thất bại sẽ ngắt luồng request ngay lập tức và trả về HTTP 401 Unauthorized, đảm bảo an toàn cho các Service và cơ sở dữ liệu."
    ]
    for para in doc.paragraphs:
        if para.text.strip() == "Hình 4: Ảnh gói API và bảo mật":
            insert_text_paragraphs(doc, para, hinh4_desc)
            break

    # Add explanation for Hình 5: Ảnh gói dịch vụ và phân hệ AI
    hinh5_desc = [
        "Thuyết minh Hình 5: Service Layer chứa core logic. AIService được thiết kế theo dạng Singleton để điều phối toàn bộ các yêu cầu sinh sinh học (WBS, Subtasks, Story Points). Nó sử dụng Interface IAIServiceProvider để đóng gói logic tương tác với LLM API. OpenRouterGatewayProvider thực thi giao diện này, tích hợp cùng TaskRoutingMap để xác định cấu hình model (ví dụ: gemini-2.5-pro cho code/estimation và claude-3.5-sonnet cho decomposition)."
    ]
    for para in doc.paragraphs:
        if para.text.strip() == "Hình 5: Ảnh gói dịch vụ và phân hệ AI":
            insert_text_paragraphs(doc, para, hinh5_desc)
            break

    # Add explanation for Hình 6: Ảnh gói ước lượng RAG và tích hợp AI
    hinh6_desc = [
        "Thuyết minh Hình 6: Phân hệ ước lượng RAG và đồng bộ hóa bao gồm EstimationService liên kết với QdrantService để sinh vector embedding và truy vấn các task tương tự. JiraExportService chuyển hóa các WBS items thành định dạng tài liệu Atlassian (ADF) thông qua ADFConverter trước khi đẩy lên Jira Cloud REST API. GitHubExportService xử lý authentication qua OAuth và đồng bộ issues trực tiếp."
    ]
    for para in doc.paragraphs:
        if para.text.strip() == "Hình 6: Ảnh gói ước lượng RAG và tích hợp AI":
            insert_text_paragraphs(doc, para, hinh6_desc)
            break

    # ----------------------------------------------------
    # TASK 2: Database Design (MongoDB Schema & PlantUML ERD)
    # ----------------------------------------------------
    print("Applying Task 2: Database Design (MongoDB / Qdrant)...")
    
    db_intro = [
        "Hệ thống APMP lưu trữ dữ liệu nghiệp vụ chính tại cơ sở dữ liệu MongoDB thông qua thư viện Mongoose ORM và dữ liệu vector phục vụ RAG tại Qdrant Vector Store.",
        "Dưới đây là sơ đồ thực thể liên kết (ERD) đặc tả mối quan hệ giữa các thực thể trong MongoDB:"
    ]
    
    erd_plantuml = [
        "@startuml APMP_ERD",
        "!theme plain",
        "skinparam classAttributeIconSize 0",
        "",
        "entity User {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * email : String <<Unique>>",
        "  * passwordHash : String",
        "  name : String",
        "  createdAt : Date",
        "  updatedAt : Date",
        "}",
        "",
        "entity UserIntegration {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * userId : ObjectId <<FK>>",
        "  * platform : String (\"jira\" | \"github\")",
        "  * accessToken : String (Encrypted)",
        "  refreshToken : String (Encrypted)",
        "  tokenExpiresAt : Date",
        "  jiraCloudId : String",
        "  githubUsername : String",
        "  createdAt : Date",
        "  updatedAt : Date",
        "}",
        "",
        "entity Project {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * title : String",
        "  * userId : ObjectId <<FK>>",
        "  * status : String (\"active\" | \"archived\" | \"completed\")",
        "  githubRepo : String",
        "  jiraProjectKey : String",
        "  createdAt : Date",
        "  updatedAt : Date",
        "}",
        "",
        "entity SourceOfTruth {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * projectId : ObjectId <<FK, Unique>>",
        "  content : String",
        "  * versionNumber : Number",
        "  versionHistory : Array<IVersionSnapshot>",
        "  createdAt : Date",
        "  updatedAt : Date",
        "}",
        "",
        "entity WBSConfig {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * projectId : ObjectId <<FK, Unique>>",
        "  methodology : String (\"scrum\" | \"kanban\" | \"waterfall\")",
        "  techStack : String",
        "  teamSize : Number",
        "  createdAt : Date",
        "  updatedAt : Date",
        "}",
        "",
        "entity WBSItem {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * projectId : ObjectId <<FK>>",
        "  parentId : ObjectId <<FK>>",
        "  sourceOfTruthId : ObjectId <<FK>>",
        "  * title : String",
        "  description : String",
        "  * type : String (\"epic\" | \"feature\" | \"story\" | \"task\" | \"subtask\")",
        "  * status : String (\"ai_generated\" | \"reviewed\" | \"approved\" | \"rejected\")",
        "  methodology : String",
        "  acceptanceCriteria : Array<String>",
        "  sourceRequirements : Array<String>",
        "  order : Number",
        "  aiGenerated : Boolean",
        "  createdAt : Date",
        "  updatedAt : Date",
        "}",
        "",
        "entity StoryPoint {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * wbsItemId : ObjectId <<FK, Unique>>",
        "  * projectId : ObjectId <<FK>>",
        "  ragReferences : Array<RAGReference>",
        "  aiSuggestedPoints : Number",
        "  finalPoints : Number",
        "  rationale : String",
        "  confidence : Number",
        "  decidedBy : ObjectId <<FK>>",
        "  createdAt : Date",
        "}",
        "",
        "entity ExternalSync {",
        "  * _id : ObjectId <<PK>>",
        "  --",
        "  * projectId : ObjectId <<FK>>",
        "  * wbsItemId : ObjectId <<FK>>",
        "  * platform : String (\"jira\" | \"github\")",
        "  * externalId : String",
        "  externalUrl : String",
        "  syncStatus : String (\"pending\" | \"synced\" | \"failed\" | \"conflict\")",
        "  errorMessage : String",
        "  lastSyncedAt : Date",
        "  createdAt : Date",
        "}",
        "",
        "User \"1\" *-- \"0..*\" UserIntegration : \"has\"",
        "User \"1\" o-- \"0..*\" Project : \"owns\"",
        "Project \"1\" *-- \"1\" SourceOfTruth : \"has\"",
        "Project \"1\" *-- \"1\" WBSConfig : \"configured by\"",
        "Project \"1\" o-- \"0..*\" WBSItem : \"contains\"",
        "WBSItem \"1\" *-- \"0..1\" WBSItem : \"parent-child hierarchy (parentId)\"",
        "WBSItem \"1\" *-- \"0..1\" StoryPoint : \"estimated by\"",
        "WBSItem \"1\" *-- \"0..*\" ExternalSync : \"tracked in\"",
        "StoryPoint \"0..*\" o-- \"1\" User : \"decidedBy\"",
        "@enduml"
    ]
    
    db_schemas_desc = [
        "Đặc tả chi tiết các Collection chính trong MongoDB:",
        "",
        "1. Collection: users",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- email: String (Bắt buộc, duy nhất) - Email đăng nhập",
        "- passwordHash: String (Bắt buộc) - Mật khẩu đã mã hoá bcrypt",
        "- name: String (Tùy chọn) - Tên người dùng",
        "- timestamps: true (createdAt, updatedAt)",
        "",
        "2. Collection: projects",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- title: String (Bắt buộc) - Tiêu đề dự án",
        "- userId: Schema.Types.ObjectId (Bắt buộc, ref: 'User') - Chủ sở hữu",
        "- status: String (enum: ['active', 'archived', 'completed']) - Trạng thái dự án",
        "- githubRepo: String (Tùy chọn) - Đường dẫn kho Github (ví dụ: 'owner/repo')",
        "- jiraProjectKey: String (Tùy chọn) - Mã dự án Jira (ví dụ: 'APMP')",
        "- timestamps: true (createdAt, updatedAt)",
        "",
        "3. Collection: sourceoftruths",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- projectId: Schema.Types.ObjectId (Bắt buộc, ref: 'Project', duy nhất)",
        "- content: String (Tùy chọn) - Nội dung yêu cầu (SRS text)",
        "- versionNumber: Number (Bắt buộc, mặc định: 1) - Phiên bản hiện tại",
        "- versionHistory: [{ versionNumber: Number, content: String, savedAt: Date }] - Lịch sử snapshot",
        "- timestamps: true (createdAt, updatedAt)",
        "",
        "4. Collection: wbsconfigs",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- projectId: Schema.Types.ObjectId (Bắt buộc, ref: 'Project', duy nhất)",
        "- methodology: String (enum: ['scrum', 'kanban', 'waterfall']) - Quy trình phát triển",
        "- techStack: String (Tùy chọn) - Công nghệ sử dụng",
        "- teamSize: Number (Tùy chọn) - Số lượng thành viên",
        "- timestamps: true (createdAt, updatedAt)",
        "",
        "5. Collection: wbsitems",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- projectId: Schema.Types.ObjectId (Bắt buộc, ref: 'Project', index)",
        "- parentId: Schema.Types.ObjectId (Tùy chọn, ref: 'WBSItem', index) - Link đến node cha",
        "- sourceOfTruthId: Schema.Types.ObjectId (Tùy chọn, ref: 'SourceOfTruth') - Link đến tài liệu gốc",
        "- title: String (Bắt buộc) - Tên hạng mục",
        "- description: String (Tùy chọn) - Mô tả chi tiết",
        "- type: String (enum: ['epic', 'feature', 'story', 'task', 'subtask']) - Cấp độ phân rã",
        "- status: String (enum: ['ai_generated', 'reviewed', 'approved', 'rejected'])",
        "- methodology: String (enum: ['scrum', 'kanban', 'waterfall'])",
        "- acceptanceCriteria: [String] - Tiêu chí nghiệm thu",
        "- sourceRequirements: [String] - ID của các phân đoạn SRS liên quan",
        "- order: Number (Mặc định: 0) - Thứ tự hiển thị",
        "- aiGenerated: Boolean (Mặc định: true) - Gắn cờ do AI sinh ra",
        "- timestamps: true (createdAt, updatedAt)",
        "",
        "6. Collection: storypoints (Lưu trong bảng Estimation)",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- wbsItemId: Schema.Types.ObjectId (Bắt buộc, ref: 'WBSItem', duy nhất)",
        "- projectId: Schema.Types.ObjectId (Bắt buộc, ref: 'Project')",
        "- ragReferences: [{ similarProjectId: Schema.Types.ObjectId, similarItemTitle: String, similarItemPoints: Number, similarityScore: Number }]",
        "- aiSuggestedPoints: Number - Điểm do AI gợi ý",
        "- finalPoints: Number - Điểm chốt cuối cùng (do người dùng chỉnh sửa)",
        "- rationale: String - Lý do ước lượng điểm",
        "- confidence: Number (từ 0 đến 1) - Độ tin cậy của gợi ý AI",
        "- decidedBy: Schema.Types.ObjectId (Tùy chọn, ref: 'User') - Người chốt điểm",
        "- createdAt: Date (Bắt buộc, tự động sinh)",
        "",
        "7. Collection: externalsyncs",
        "- _id: Schema.Types.ObjectId (Khóa chính)",
        "- projectId: Schema.Types.ObjectId (Bắt buộc, ref: 'Project')",
        "- wbsItemId: Schema.Types.ObjectId (Bắt buộc, ref: 'WBSItem', unique cùng platform)",
        "- platform: String (enum: ['jira', 'github'])",
        "- externalId: String (Bắt buộc) - Key bên ngoài (ví dụ: 'JIRA-123' hoặc issue number)",
        "- externalUrl: String (Tùy chọn) - Link đến issue trang ngoài",
        "- syncStatus: String (enum: ['pending', 'synced', 'failed', 'conflict'])",
        "- errorMessage: String (Tùy chọn) - Chi tiết lỗi đồng bộ",
        "- lastSyncedAt: Date (Tùy chọn)",
        "- createdAt: Date",
        "",
        "Đặc tả Cơ sở Dữ liệu Vector Qdrant:",
        "Qdrant lưu trữ các vector đại diện cho lịch sử các issues để thực hiện cơ chế RAG (Retrieval-Augmented Generation).",
        "- Collection: issues",
        "- Kích thước Vector (Dimension): 768 chiều (sinh từ model: google/gemini-embedding-2)",
        "- Độ đo khoảng cách (Distance Metric): Cosine",
        "- Payload Schema chứa các trường:",
        "  + issuekey: String - Mã định danh issue",
        "  + idproject: String - Mã dự án gốc",
        "  + title: String - Tiêu đề của Task/Issue",
        "  + description: String - Mô tả chi tiết của Task/Issue",
        "  + storypoints: Number - Số điểm story points thực tế"
    ]

    # Replace the placeholders under "Thiết kế cơ sở dữ liệu"
    # The paragraph 108 in the word doc contains guidelines.
    for para in doc.paragraphs:
        if "Phần này có độ dài từ hai đến bốn trang. Sinh viên thiết kế, vẽ và giải thích biểu đồ thực thể liên kết" in para.text:
            para.text = "Sơ đồ thực thể liên kết ERD (PlantUML) đặc tả cơ sở dữ liệu hệ thống APMP:"
            # Insert the UML code block
            curr = add_code_block(doc, para, erd_plantuml)
            # Insert the introduction and schemas desc after it
            insert_text_paragraphs(doc, curr, db_schemas_desc)
            break
            
    # Clean up the next paragraph which is the example instruction (Para 109)
    for para in doc.paragraphs:
        if "Ví dụ: Xác định các thực thể dữ liệu, liên kết và thuộc tính" in para.text:
            para.text = "" # Clear it
            break

    # ----------------------------------------------------
    # TASK 3: Class & Workflow Design (Class/Sequence PlantUML)
    # ----------------------------------------------------
    print("Applying Task 3: Class & Workflow Design...")
    
    class_intro = [
        "Sơ đồ lớp kiến trúc (Class Diagram) đặc tả thuộc tính và phương thức chính của các Service nghiệp vụ và AI Providers:",
    ]
    
    class_plantuml = [
        "@startuml APMP_Class_Diagram",
        "!theme plain",
        "skinparam classAttributeIconSize 0",
        "",
        "class useProjectStore {",
        "  + currentProject : ProjectDTO",
        "  + wbsItems : List<WBSItemDTO>",
        "  + loadProject(id : String) : void",
        "  + generateWBS(config : Object) : void",
        "  + estimateStoryPoints(itemId : String) : void",
        "}",
        "",
        "class APIClient {",
        "  + get(url : String) : Promise<Response>",
        "  + post(url : String, body : Object) : Promise<Response>",
        "}",
        "",
        "class RouteHandler {",
        "  + GET(req : Request) : Promise<Response>",
        "  + POST(req : Request) : Promise<Response>",
        "}",
        "",
        "class AuthMiddleware {",
        "  + verifyJWT(req : Request) : Promise<UserDTO>",
        "}",
        "",
        "class ProjectService {",
        "  + createProject(input : CreateProjectInput) : Promise<ProjectDTO>",
        "  + findProjectById(id : String) : Promise<ProjectDTO>",
        "}",
        "",
        "class WBSItemService {",
        "  + generateWBS(projectId : String, config : Object) : Promise<WBSItemDTO[]>",
        "  + findWBSItemsByProjectId(projectId : String) : Promise<WBSItemDTO[]>",
        "  + generateDeveloperSubtasks(taskId : String) : Promise<WBSItemDTO[]>",
        "}",
        "",
        "class EstimationService {",
        "  + estimateStoryPoints(wbsItemId : String, userId : String) : Promise<StoryPointDTO>",
        "  + getStoryPointEstimation(wbsItemId : String) : Promise<StoryPointDTO>",
        "  + updateStoryPointEstimate(wbsItemId : String, points : Number) : Promise<StoryPointDTO>",
        "}",
        "",
        "class QdrantService {",
        "  + generateEmbedding(text : String) : Promise<Number[]>",
        "  + searchSimilarIssues(text : String, limit : Number) : Promise<Object[]>",
        "}",
        "",
        "class AIService <<Singleton>> {",
        "  - {static} instance : AIService",
        "  - provider : IAIServiceProvider",
        "  + {static} getInstance() : AIService",
        "  + generateWBS(sot : String, config : Object) : Promise<Object[]>",
        "  + estimateStoryPoints(title : String, desc : String, refs : Object[]) : Promise<Object>",
        "}",
        "",
        "interface IAIServiceProvider {",
        "  + generateWBS(sot : String, config : Object) : Promise<Object[]>",
        "  + estimateStoryPoints(title : String, desc : String, refs : Object[]) : Promise<Object>",
        "}",
        "",
        "class OpenRouterGatewayProvider {",
        "  - apiKey : String",
        "  - routingMap : TaskRoutingMap",
        "  + generateWBS(sot : String, config : Object) : Promise<Object[]>",
        "  + estimateStoryPoints(title : String, desc : String, refs : Object[]) : Promise<Object>",
        "}",
        "",
        "useProjectStore ..> APIClient : calls",
        "APIClient ..> RouteHandler : REST",
        "RouteHandler ..> AuthMiddleware : authenticates",
        "RouteHandler ..> ProjectService : delegates",
        "RouteHandler ..> WBSItemService : delegates",
        "RouteHandler ..> EstimationService : delegates",
        "WBSItemService ..> AIService : AI",
        "EstimationService ..> AIService : AI",
        "EstimationService ..> QdrantService : Semantic Search",
        "AIService o-- IAIServiceProvider : delegates",
        "IAIServiceProvider <|.. OpenRouterGatewayProvider : implements",
        "@enduml"
    ]
    
    workflows_desc = [
        "",
        "Biểu đồ trình tự (Sequence Diagrams) đặc tả luồng truyền điệp cho 4 nghiệp vụ cốt lõi:",
        "",
        "Nghiệp vụ 1: Upload và Phân tích Yêu cầu SRS",
    ]
    
    seq1_puml = [
        "@startuml Sequence_SRS_Ingestion",
        "actor User as U",
        "participant \"ProjectDashboard (UI)\" as UI",
        "participant \"RouteHandler (/api/files)\" as API",
        "participant \"AuthMiddleware\" as Auth",
        "participant \"FileService\" as FS",
        "participant \"pdfExtractor\" as PE",
        "participant \"SourceOfTruthService\" as SOTS",
        "database MongoDB as DB",
        "",
        "U -> UI: Upload SRS file (.docx/.pdf)",
        "UI -> API: POST /api/files (multipart/form-data)",
        "API -> Auth: verifyJWT(request)",
        "Auth --> API: Decoded JWT (User Context)",
        "API -> FS: saveFile(buffer, filename)",
        "FS -> DB: Create File document",
        "DB --> FS: File document saved",
        "FS -> PE: extractText(fileStream)",
        "PE --> FS: Raw text content",
        "FS -> SOTS: updateOrCreateSOT(projectId, extractedText)",
        "SOTS -> DB: Find existing SOT",
        "alt SOT exists",
        "  SOTS -> DB: snapshotVersion()",
        "  SOTS -> DB: Update content & versionNumber + 1",
        "else SOT new",
        "  SOTS -> DB: Create new SourceOfTruth",
        "end",
        "DB --> SOTS: SOT saved",
        "SOTS --> FS: SOT Document",
        "FS --> API: Ingestion result",
        "API --> UI: HTTP 200 OK (SOT metadata & text preview)",
        "UI --> U: Display requirements text on panel",
        "@enduml"
    ]
    
    workflows_desc2 = [
        "",
        "Nghiệp vụ 2: Phân rã WBS tự động bằng AI",
    ]
    
    seq2_puml = [
        "@startuml Sequence_WBS_Generation",
        "actor User as U",
        "participant \"WBSPanel (UI)\" as UI",
        "participant \"RouteHandler (/api/wbs/[id])\" as API",
        "participant \"WBSItemService\" as WIS",
        "participant \"SourceOfTruthService\" as SOTS",
        "participant \"AIService\" as AIS",
        "participant \"OpenRouterGatewayProvider\" as ORP",
        "database MongoDB as DB",
        "",
        "U -> UI: Click \"Generate WBS\"",
        "UI -> API: POST /api/wbs/[id] { methodology, techStack }",
        "API -> WIS: generateWBSForProject(projectId, config)",
        "WIS -> SOTS: getSourceOfTruth(projectId)",
        "SOTS -> DB: Query SOT content",
        "DB --> SOTS: SOT raw text",
        "SOTS --> WIS: Return SOT content",
        "WIS -> AIS: generateWBS(sotText, config)",
        "AIS -> ORP: generateWBS(sotText, config)",
        "ORP -> \"OpenRouter API\": POST /chat/completions (model: anthropic/claude-3-sonnet)",
        "\"OpenRouter API\" --> ORP: JSON response (Epics, Stories, Tasks)",
        "ORP --> AIS: Parsed JSON array",
        "AIS --> WIS: Tree of WBS items",
        "WIS -> DB: Delete existing WBS items for project",
        "WIS -> DB: Bulk write new WBSItems (Level 1 to 4)",
        "DB --> WIS: Bulk write success",
        "WIS --> API: Generated WBS item array",
        "API --> UI: HTTP 200 OK (WBS tree list)",
        "UI --> U: Render WBS tree on UI",
        "@enduml"
    ]
    
    workflows_desc3 = [
        "",
        "Nghiệp vụ 3: Ước lượng Story Points bằng kỹ thuật RAG",
    ]
    
    seq3_puml = [
        "@startuml Sequence_RAG_Estimation",
        "actor User as U",
        "participant \"WBSPanel (UI)\" as UI",
        "participant \"RouteHandler (/api/wbs/[id]/estimate)\" as API",
        "participant \"EstimationService\" as ES",
        "participant \"QdrantService\" as QS",
        "participant \"AIService\" as AIS",
        "database Qdrant as QD",
        "database MongoDB as DB",
        "",
        "U -> UI: Click \"Estimate Task\"",
        "UI -> API: POST /api/wbs/[id]/estimate { wbsItemId }",
        "API -> ES: estimateStoryPoints(wbsItemId, userId)",
        "ES -> DB: Find WBSItem by ID",
        "DB --> ES: Return WBSItem (Title, Description)",
        "ES -> QS: searchSimilarIssues(queryText, limit=3)",
        "QS -> QS: generateEmbedding(queryText)",
        "QS -> \"OpenRouter Embeddings API\": POST /embeddings (google/gemini-embedding-2)",
        "\"OpenRouter Embeddings API\" --> QS: 768d vector",
        "QS -> QD: POST /collections/issues/points/search { vector }",
        "QD --> QS: Top 3 similar issues (Title, Desc, StoryPoints, Score)",
        "QS --> ES: Array of matched analogues",
        "ES -> AIS: estimateStoryPoints(title, desc, references)",
        "AIS -> \"OpenRouter API\": POST /chat/completions (model: google/gemini-2.5-pro)",
        "\"OpenRouter API\" --> AIS: Suggested Points & Rationale JSON",
        "AIS --> ES: Return points, rationale, confidence",
        "ES -> DB: Create/Update StoryPoint document",
        "DB --> ES: StoryPoint saved",
        "ES --> API: StoryPoint estimation DTO",
        "API --> UI: HTTP 200 OK (Estimation result)",
        "UI --> U: Open modal displaying points, rationale, and RAG references",
        "@enduml"
    ]
    
    workflows_desc4 = [
        "",
        "Nghiệp vụ 4: Đồng bộ Backlog sang Jira Cloud",
    ]
    
    seq4_puml = [
        "@startuml Sequence_Backlog_Sync",
        "actor User as U",
        "participant \"IntegrationsPanel (UI)\" as UI",
        "participant \"RouteHandler (/api/jira/export)\" as API",
        "participant \"JiraExportService\" as JES",
        "participant \"JiraOAuthService\" as JOS",
        "participant \"ADFConverter\" as ADF",
        "database MongoDB as DB",
        "",
        "U -> UI: Click \"Sync to Jira\"",
        "UI -> API: POST /api/jira/export { projectId }",
        "API -> JES: exportProjectBacklog(projectId, userId)",
        "JES -> DB: Get WBSItems and StoryPoints for project",
        "DB --> JES: List of WBS items with estimation",
        "JES -> JOS: getAccessToken(userId)",
        "JOS -> DB: Find UserIntegration (Jira OAuth tokens)",
        "DB --> JOS: Encrypted accessToken & refreshToken",
        "JOS -> JOS: Decrypt token using AES-GCM",
        "JOS --> JES: Clear accessToken",
        "loop For each WBS item (Epic -> Story -> Task)",
        "  JES -> ADF: convertToADF(description)",
        "  ADF --> JES: Atlassian Document Format JSON",
        "  JES -> \"Jira Cloud API\": POST /rest/api/3/issue { fields: summary, description, customfield_SP }",
        "  \"Jira Cloud API\" --> JES: Issue response { id, key, self }",
        "  JES -> DB: Save ExternalSync record",
        "  DB --> JES: Saved",
        "end",
        "JES --> API: Export status summary",
        "API --> UI: HTTP 200 OK (Sync complete)",
        "UI --> U: Show \"Synced\" status badges next to items",
        "@enduml"
    ]
    
    # Replace "Thiết kế lớp" instruction text
    for para in doc.paragraphs:
        if "Sinh viên trình bày thiết kế chi tiết các thuộc tính và phương thức" in para.text:
            para.text = "Mô hình thiết kế chi tiết lớp và mối quan hệ giữa các cấu phần trong hệ thống APMP:"
            curr = add_code_block(doc, para, class_plantuml)
            curr = insert_text_paragraphs(doc, curr, workflows_desc)
            curr = add_code_block(doc, curr, seq1_puml)
            curr = insert_text_paragraphs(doc, curr, workflows_desc2)
            curr = add_code_block(doc, curr, seq2_puml)
            curr = insert_text_paragraphs(doc, curr, workflows_desc3)
            curr = add_code_block(doc, curr, seq3_puml)
            curr = insert_text_paragraphs(doc, curr, workflows_desc4)
            curr = add_code_block(doc, curr, seq4_puml)
            break
            
    # Clean the second instruction paragraph under class design
    for para in doc.paragraphs:
        if "Để minh họa thiết kế lớp, sinh viên thiết kế luồng truyền thông điệp" in para.text:
            para.text = ""
            break

    # ----------------------------------------------------
    # TASK 4: REST API Specifications & GUI Design
    # ----------------------------------------------------
    print("Applying Task 4: REST API Specifications...")
    
    api_specs = [
        "Danh sách đặc tả các API chính phía Backend của hệ thống APMP:",
        "",
        "1. API: POST /api/auth/register",
        "- Mô tả: Cho phép người dùng đăng ký tài khoản mới.",
        "- Auth: Không bắt buộc (Public)",
        "- Request Input (JSON):",
        "  + email (String, bắt buộc) - Email hợp lệ",
        "  + password (String, bắt buộc) - Mật khẩu tối thiểu 6 ký tự",
        "  + name (String, tùy chọn) - Tên hiển thị",
        "- Response Output:",
        "  + 201 Created: { success: true, user: { id, email, name } }",
        "  + 400 Bad Request: { success: false, error: 'Email already exists' }",
        "",
        "2. API: POST /api/auth/login",
        "- Mô tả: Xác thực tài khoản người dùng và thiết lập session cookie.",
        "- Auth: Không bắt buộc (Public)",
        "- Request Input (JSON):",
        "  + email (String, bắt buộc)",
        "  + password (String, bắt buộc)",
        "- Response Output:",
        "  + 200 OK: { success: true, user: { id, email, name } } (Đặt JWT trong HttpOnly cookie 'token')",
        "  + 401 Unauthorized: { success: false, error: 'Invalid credentials' }",
        "",
        "3. API: POST /api/projects",
        "- Mô tả: Tạo một dự án quản lý mới gắn với người dùng hiện tại.",
        "- Auth: Bắt buộc (JWT)",
        "- Request Input (JSON):",
        "  + title (String, bắt buộc) - Tên dự án",
        "  + status (String, tùy chọn, mặc định: 'active')",
        "  + githubRepo (String, tùy chọn) - định dạng 'owner/repo'",
        "  + jiraProjectKey (String, tùy chọn)",
        "- Response Output:",
        "  + 201 Created: { success: true, project: { id, title, userId, status, githubRepo, jiraProjectKey, createdAt } }",
        "",
        "4. API: POST /api/files",
        "- Mô tả: Tải lên tài liệu SRS (.pdf/.docx) để hệ thống tự động bóc tách text và lưu thành Source of Truth.",
        "- Auth: Bắt buộc (JWT)",
        "- Request Input (Multipart Form-Data):",
        "  + file (File, bắt buộc) - Định dạng PDF hoặc DOCX, tối đa 10MB",
        "  + projectId (String, bắt buộc) - ID của dự án đích",
        "- Response Output:",
        "  + 200 OK: { success: true, fileId: '...', textLength: 12500 }",
        "  + 400 Bad Request: { success: false, error: 'Invalid file format' }",
        "",
        "5. API: POST /api/wbs/[id]",
        "- Mô tả: Gọi phân hệ AI để tự động phân rã văn bản Source of Truth thành cây WBS 4 cấp.",
        "- Auth: Bắt buộc (JWT)",
        "- Path Parameter: id - ID của dự án",
        "- Request Input (JSON):",
        "  + methodology (String, tùy chọn) - 'scrum' | 'kanban' | 'waterfall'",
        "  + techStack (String, tùy chọn) - Mô tả techstack",
        "- Response Output:",
        "  + 200 OK: { success: true, itemsCount: 45, items: [ { id, title, parentId, type, status, order } ] }",
        "",
        "6. API: POST /api/wbs/[id]/estimate",
        "- Mô tả: Ước lượng story points cho 1 hạng mục công việc (Task/Story) dựa trên kỹ thuật RAG kết hợp LLM.",
        "- Auth: Bắt buộc (JWT)",
        "- Request Input (JSON):",
        "  + wbsItemId (String, bắt buộc) - ID của WBS Item cần ước lượng",
        "- Response Output:",
        "  + 200 OK: { success: true, estimate: { id, wbsItemId, aiSuggestedPoints, finalPoints, rationale, confidence } }",
        "",
        "7. API: POST /api/jira/export",
        "- Mô tả: Đồng bộ toàn bộ backlog/WBS đã phê duyệt lên dự án Jira của người dùng.",
        "- Auth: Bắt buộc (JWT)",
        "- Request Input (JSON):",
        "  + projectId (String, bắt buộc) - ID dự án APMP cần đồng bộ",
        "- Response Output:",
        "  + 200 OK: { success: true, syncedCount: 15, issues: [ { wbsItemId, externalId: 'APMP-42', externalUrl: '...' } ] }"
    ]
    
    gui_specs = [
        "Đặc tả Giao diện Người dùng (GUI):",
        "",
        "1. Các Thông số Giao diện Kỹ thuật:",
        "- Độ phân giải tối ưu: 1440x900px trở lên cho phiên bản Desktop, hiển thị dạng Single-page Application (SPA).",
        "- Hệ màu sắc (Palette): Sử dụng cấu hình màu sắc tương phản cao (Dark Mode làm chủ đạo):",
        "  + Màu nền chính (Background): HSL Slate-900 (#0f172a)",
        "  + Màu nền panel phụ: HSL Slate-800 (#1e293b)",
        "  + Màu nhấn (Accent color): Emerald-500 (#10b981) - chỉ thị trạng thái thành công/phê duyệt, Violet-600 (#7c3aed) - chỉ thị thương hiệu và tương tác chính.",
        "  + Màu chữ: HSL Slate-100 (#f1f5f9) cho tiêu đề và Slate-400 (#94a3b8) cho nội dung mô tả.",
        "- Typography: Sử dụng font chữ hệ thống Inter làm chủ đạo cho giao diện và font Geist Mono làm font hiển thị code/cấu trúc dữ liệu.",
        "",
        "2. Cấu trúc Bố cục Giao diện (Layout):",
        "- Hệ thống bố cục chia làm 3 khu vực chính:",
        "  + Sidebar điều hướng (Left-side bar): Danh sách dự án hiện tại, trạng thái kết nối tích hợp (Jira/GitHub), và nút cấu hình cá nhân.",
        "  + Khu vực Workspace chính (Central Workspace): Hiển thị dưới dạng các Tab tương ứng với các phân hệ: Source of Truth (Xem và sửa tài liệu gốc), Conflicts (Bảng so sánh và duyệt mâu thuẫn yêu cầu), và WBS Editor (Dưới dạng một Tree-Grid bảng phân cấp cho phép xem/sửa, chấm điểm Story Points, đồng bộ).",
        "  + Drawer chi tiết ước lượng (Right-side Drawer): Hiển thị chi tiết lý do ước lượng Story Point từ AI, danh sách 3 issue tương đồng tìm được từ Qdrant Vector DB kèm độ tương đồng (similarity score).",
        "",
        "Lưu ý: Sinh viên cần dán bổ sung các ảnh chụp màn hình GUI thực tế (màn hình Dashboard, màn hình thiết lập tích hợp, màn hình WBS Editor và Modal chi tiết ước lượng) tại đây."
    ]

    # Find where API design starts. Paragraph 119 starts the API examples.
    # We will search for basic URL example: "Đường dẫn cơ bản"
    api_target_para = None
    for para in doc.paragraphs:
        if "Đường dẫn cơ bản" in para.text or "https://ABC.def/it4788/" in para.text:
            api_target_para = para
            break
            
    if api_target_para:
        api_target_para.text = "" # Clear it
        curr = insert_text_paragraphs(doc, api_target_para, api_specs)
        
        # We also want to clean up the rest of the template paragraphs for signup/signin
        # which are from Paragraph 120 to 136 in the original file.
        # Let's search and clear paragraphs that mention "API: /signup" or "phonenumber"
        for para in doc.paragraphs:
            txt = para.text
            if any(term in txt for term in ["API: /signup", "phonenumber", "Kết quả: 1000|OK", "Chi tiết input / output"]):
                para.text = ""
    else:
        print("Warning: API placeholder target paragraph not found by matching 'Đường dẫn cơ bản'")

    # Replace GUI instruction placeholders.
    gui_target_para = None
    for para in doc.paragraphs:
        if "Sinh viên đặc tả thông tin về màn hình mà ứng dụng của mình hướng tới" in para.text:
            gui_target_para = para
            break
            
    if gui_target_para:
        gui_target_para.text = ""
        insert_text_paragraphs(doc, gui_target_para, gui_specs)
        
        # Clean up next paragraphs (139, 140)
        for para in doc.paragraphs:
            if any(term in para.text for term in ["Sinh viên xây dựng các thiết kế giao diện mockup", "Sau cùng sinh viên đưa ra một số hình ảnh minh họa"]):
                para.text = ""
    else:
        print("Warning: GUI placeholder target paragraph not found")

    # Save the updated document
    out_path = os.path.join("docs", "Tai lieu thiet ke phan mem SDD (1).docx")
    print(f"Saving final changes to: {out_path} ...")
    doc.save(out_path)
    print("SDD Document completed successfully!")

if __name__ == "__main__":
    main()
