import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys
import re

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def insert_para_after(doc, target_para, text, bold=False, italic=False, font_size_pt=12, alignment=0, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = docx.shared.Pt(space_before)
    p.paragraph_format.space_after = docx.shared.Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = alignment
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    
    target_para._p.addnext(p._p)
    return p

def main():
    doc_path = os.path.join("docs", "Mẫu báo cáo cuối kỳ Project, GR.docx")
    if not os.path.exists(doc_path):
        print(f"Error: Could not find document at {doc_path}")
        sys.exit(1)
        
    print(f"Loading document: {doc_path} ...")
    doc = docx.Document(doc_path)
    
    # 1. Locate Chapter 6 boundaries
    start_idx = None
    end_idx = None
    
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip().lower().replace('\xa0', ' ').replace('\u200b', '')
        if txt == "kết luận và hướng phát triển" and idx > 350:
            start_idx = idx
        elif txt == "danh mục thuật ngữ" or txt == "tài liệu tham khảo" or (txt.startswith("tài liệu") and "tham khảo" in txt) or idx == len(doc.paragraphs) - 1:
            end_idx = idx
            # We don't break immediately to find the closest one after start_idx
            if start_idx is not None:
                break
                
    if start_idx is None or end_idx is None:
        print(f"Error: Could not find Chapter 6 boundaries (start: {start_idx}, end: {end_idx})")
        sys.exit(1)
        
    print(f"Chapter 6 starts at index {start_idx} ('{doc.paragraphs[start_idx].text.strip()}') and ends before {end_idx} ('{doc.paragraphs[end_idx].text.strip()}')")
    
    # Rename title
    doc.paragraphs[start_idx].text = "Chương 6: Kết luận và hướng phát triển"
    title_run = doc.paragraphs[start_idx].runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = docx.shared.Pt(16)
    title_run.bold = True
    
    # Delete placeholder paragraphs
    paras_to_delete = [doc.paragraphs[k] for k in range(start_idx + 1, end_idx)]
    for p in paras_to_delete:
        p_element = p._p
        p_element.getparent().remove(p_element)
    print(f"Removed {len(paras_to_delete)} placeholder paragraphs.")
    
    # Save and reload
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    
    # Re-locate Chapter 6 title index
    start_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if "chương 6: kết luận và hướng phát triển" in para.text.strip().lower():
            start_idx = idx
            break
            
    cursor = doc.paragraphs[start_idx]
    
    # 2. Define refined Chapter 6 content paragraphs
    ch6_sections = [
        ("6.1. Kết luận", True),
        (
            "Trong suốt quá trình thực hiện đồ án tốt nghiệp, nhóm nghiên cứu đã đạt được các kết quả quan trọng sau: (1) Xây dựng thành công hệ thống APMP hỗ trợ vai trò BA và PM trong giai đoạn Inception & Planning; (2) Hiện thực hóa quy trình trích xuất tài liệu SRS đa phương thức và tự động đối chiếu phát hiện mâu thuẫn chéo giữa văn bản thô và hình vẽ sơ đồ bằng AI; (3) Thực hiện phân rã WBS 4 cấp tự động theo cấu hình quy trình của từng dự án; (4) Triển khai giải pháp ước lượng Story Points khách quan bằng kỹ thuật RAG, kết nối kho tri thức dự án tương đồng lịch sử thông qua Qdrant Vector DB; (5) Hỗ trợ đồng bộ hóa backlog nhanh chóng 1-click sang Jira Cloud và GitHub.",
            False
        ),
        (
            "Bên cạnh các kết quả đạt được, đồ án vẫn tồn tại một số hạn chế nhất định cần khắc phục: (1) Hệ thống mới chỉ hỗ trợ tốt các tệp PDF nguyên bản tạo từ trình soạn thảo (digital-born PDF), chưa tối ưu hóa hiệu năng bóc tách cho các tài liệu quét scan hoặc ảnh chụp chất lượng kém; (2) Chưa hỗ trợ không gian làm việc cộng tác đồng thời cho nhiều người dùng (Collaborative Multi-user Workspace) mà mới chỉ hỗ trợ lưu trữ trạng thái đơn lẻ.",
            False
        ),
        (
            "Thông qua quá trình nghiên cứu và thực hiện đồ án, nhóm phát triển đã rút ra những bài học kinh nghiệm sâu sắc: (1) Hiểu rõ vai trò của việc kiểm soát tham số nhiệt độ mô hình (Temperature = 0) để bảo toàn dữ liệu gốc khi làm việc với LLM; (2) Nhận thức tầm quan trọng của kỹ thuật RAG trong việc bổ sung tri thức doanh nghiệp thực tế cho mô hình AI mà không cần tốn chi phí fine-tuning tốn kém.",
            False
        ),
        
        ("6.2. Hướng phát triển", True),
        (
            "Trong tương lai, để hoàn thiện sản phẩm, nhóm nghiên cứu đề xuất nâng cấp hệ thống để tiếp nhận thêm nhiều loại tài liệu kỹ thuật chuyên sâu và chi tiết hơn ngoài tài liệu đặc tả yêu cầu SRS (như tài liệu thiết kế kiến trúc hệ thống, đặc tả cơ sở dữ liệu chi tiết, hoặc tài liệu thiết kế giao diện UI/UX), qua đó giúp mô hình LLM có đầy đủ bối cảnh kỹ thuật để phân rã các tác vụ (tasks) lập trình ở mức độ chi tiết và thực tế hơn khi đồng bộ sang Jira và GitHub. Đồng thời, hệ thống sẽ thiết lập cơ chế quản lý tài khoản và phân quyền người dùng (Role-Based Access Control) rõ ràng cho các vai trò cốt lõi như Chuyên viên Phân tích Nghiệp vụ (BA) và Quản trị Dự án (PM) để triển khai quy trình kiểm duyệt có kiểm soát (User-in-the-loop), cho phép con người trực tiếp rà soát, đánh giá tính hợp lý và hiệu chỉnh các tác vụ do AI đề xuất trước khi xuất bản. Hơn thế nữa, đồ án định hướng mở rộng phạm vi chức năng để quản lý toàn diện vòng đời dự án thay vì chỉ tập trung vào giai đoạn lập kế hoạch (Inception & Planning) như hiện tại, bằng cách phát triển thêm các module cho phép cập nhật tiến độ công việc, ghi nhận trạng thái hoàn thành của từng tác vụ và đồng bộ hai chiều thời gian thực với Jira/GitHub để trở thành một hệ quản trị dự án hoàn chỉnh.",
            False
        )
    ]
    
    # Insert paragraphs sequentially
    for text, is_header in ch6_sections:
        cursor = insert_para_after(doc, cursor, text, bold=is_header, space_before=6 if is_header else 4, space_after=6 if is_header else 4)
        
    doc.save(doc_path)
    
    # 3. Renumber all figure captions sequentially
    print("\n--- Running Auto-Renumbering of Figures ---")
    doc = docx.Document(doc_path)
    fig_pattern = re.compile(r"^Hình\s+(\d+|x|vẽ)[\s\:\-\.]*", re.IGNORECASE)
    fig_counter = 1
    
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip().replace('\xa0', ' ').replace('\u200b', '')
        match = fig_pattern.match(txt)
        if match:
            prefix = match.group(0)
            clean_text = txt[len(prefix):].strip()
            
            para.text = ""
            para.alignment = 1
            para.paragraph_format.space_before = docx.shared.Pt(3)
            para.paragraph_format.space_after = docx.shared.Pt(12)
            para.paragraph_format.line_spacing = 1.15
            
            run_bold = para.add_run(f"Hình {fig_counter}: ")
            run_bold.font.name = 'Times New Roman'
            run_bold.font.size = docx.shared.Pt(12)
            run_bold.bold = True
            
            run_normal = para.add_run(clean_text)
            run_normal.font.name = 'Times New Roman'
            run_normal.font.size = docx.shared.Pt(12)
            
            print(f"  Renumbered: Hình {fig_counter}: {clean_text[:65]}...")
            fig_counter += 1
            
    doc.save(doc_path)
    print("Chapter 6 successfully created and figures renumbered!")

if __name__ == "__main__":
    main()
